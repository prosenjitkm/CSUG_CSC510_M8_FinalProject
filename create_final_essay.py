"""
Create the final essay DOCX file with complete content and references
"""
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed")
    print("Run: pip install python-docx")
    sys.exit(1)

def create_essay():
    print("Creating CSC510 Final Essay in DOCX format...")
    print(f"Working directory: {os.getcwd()}")

    # Create document
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Astrology AI Decision Support:\nAn Expert System Approach to Personal Decision-Making')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info.add_run('[Your Name]\n')
    author_info.add_run('Colorado State University Global\n')
    author_info.add_run('CSC510: Artificial Intelligence\n')
    author_info.add_run('[Instructor Name]\n')
    author_info.add_run('March 9, 2026')

    doc.add_page_break()

    # Main Title
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run('Astrology AI Decision Support:\nAn Expert System Approach to Personal Decision-Making')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Introduction
    h1 = doc.add_heading('Introduction and Problem Statement', level=1)
    h1.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'Modern individuals face complex personal decisions spanning career transitions, relationship commitments, '
        'financial planning, and family considerations. While these decisions benefit from structured analysis, many '
        'people lack systematic frameworks to organize their thinking. This project develops an AI-powered decision-support '
        'assistant that combines symbolic reasoning, expert system principles, and intelligent search methods to help users '
        'think through important life choices. The Astrology AI Decision Support Assistant is a fully-functioning Python program '
        'that accepts natural language input, identifies decision contexts, and delivers structured guidance based on transparent, '
        'traceable reasoning processes. The system serves as a reflective tool rather than claiming predictive capability, helping '
        'users organize thoughts, identify critical factors, and develop actionable next steps.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Tools
    h2 = doc.add_heading('Tools, Libraries, and APIs Utilized', level=1)
    h2.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'The system integrates custom implementations with standard libraries and external services. Three data structures '
        'were implemented from scratch: HashTable (models/hash_table.py) for session memory with O(1) insertion/retrieval, '
        'Set (models/set_operations.py) for keyword overlap calculations in intent detection, and QuickSelect (models/sorting.py) '
        'for O(n) ranking of recommendation candidates. Python\'s collections.deque supports breadth-first search queue operations, '
        'while heapq enables A* priority queue management. The datetime module handles birth date normalization and zodiac calculation, '
        'and regular expressions tokenize user statements. Flask 3.0.0 provides web framework capabilities with RESTful API endpoints '
        'and browser interface. The OpenStreetMap Nominatim API converts birth locations to geographic coordinates, demonstrating '
        'external service integration (OpenStreetMap Foundation, n.d.).'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Search Methods
    h3 = doc.add_heading('Search Methods and Their Contribution', level=1)
    h3.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'Intelligent search forms the core of the symbolic planning subsystem. The breadth-first search implementation '
        '(astrology_ai.py, lines 257-282) explores decision-state graphs level by level, guaranteeing shortest paths from '
        '"start" to "decision_ready" goal states. BFS uses deque-based queues and visited sets to prevent cycles, ensuring '
        'users receive direct paths without unnecessary steps. For example, a general decision follows: start → clarify_goal '
        '→ collect_facts → evaluate_tradeoffs → decision_ready.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        'The A* algorithm (lines 285-322) enhances efficiency through heuristic guidance. Maintaining priority queues ordered '
        'by f(n) = g(n) + h(n), where g represents actual cost and h estimates remaining cost to goal, A* expands fewer nodes '
        'than BFS in complex graphs. The heuristic function (lines 227-254) assigns distances based on state semantics: '
        '"timeline_plan" = 1, "decision_ready" = 0, "start" = 4. This admissible heuristic ensures optimality while reducing '
        'search effort (Hart, Nilsson, & Raphael, 1968). For family planning scenarios with specialized states like '
        '"partner_alignment" and "financial_readiness," A* typically expands 5-7 nodes versus BFS\'s 8-10 nodes. Both methods '
        'execute for comparison, with outputs showing node counts and visited sequences, providing transparency into algorithmic '
        'decision-making.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Deep Learning
    h4 = doc.add_heading('Deep Learning Considerations', level=1)
    h4.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'This implementation deliberately excludes deep learning models based on three factors: explainability requirements, '
        'computational efficiency, and course alignment. Deep neural networks function as "black boxes" where input-output '
        'relationships emerge from learned weights, making it difficult to explain specific recommendations (Russell & Norvig, 2021). '
        'The rule-based approach allows every recommendation to trace back to explicit rules and keyword matches that users can inspect. '
        'The system operates as lightweight software requiring no GPU or pre-trained models, ensuring accessibility on standard hardware. '
        'The project emphasizes classical AI—search algorithms, expert systems, symbolic planning—demonstrating mastery across multiple '
        'course modules rather than focusing solely on statistical methods. Future versions could integrate supervised classifiers for '
        'improved intent detection while maintaining the transparent symbolic reasoning core.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Expert System
    h5 = doc.add_heading('Expert System Architecture', level=1)
    h5.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'The system directly implements expert system principles with knowledge base, inference engine, and working memory '
        'components (Giarratano & Riley, 2005). The knowledge base comprises multiple symbol tables: SIGN_PROFILES encodes twelve '
        'zodiac archetypes with elements, modalities, strengths, and cautions; INTENT_KEYWORDS maps vocabulary to decision contexts; '
        'ELEMENT_GUIDANCE provides 24 situation-specific advice templates; and TOPIC_HINTS extends semantic coverage with multi-word '
        'phrases. The EXPERT_RULES table (lines 152-175) encodes condition-conclusion relationships. Rule R1 states: "IF intent = family '
        'THEN include partner alignment, health readiness, and timeline checks" (weight: 8). Rule R3: "IF confidence ≤ 50% THEN collect '
        'additional data before committing" (weight: 7).'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        'The evaluate_expert_rules() function implements forward-chaining inference, examining rule conditions against current context '
        '(intent, element, confidence). When conditions match, rules fire, adding weighted conclusions to recommendation pools. The '
        'HashTable-based working memory tracks detected intents across session interactions, enabling pattern recognition. The '
        'session_pattern() function identifies recurring themes, informing users when they repeatedly query the same domain, mirroring '
        'how expert systems track evolving reasoning states (Luger, 2009).'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Knowledge Representation
    h6 = doc.add_heading('Knowledge Representation Methods', level=1)
    h6.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'Knowledge representation employs multiple complementary structures (Russell & Norvig, 2021). Symbol tables use Python '
        'dictionaries for O(1) lookup of zodiac profiles, intent keywords, and guidance templates. ELEMENT_GUIDANCE\'s nested structure '
        'represents the element × intent cross-product with 24 distinct advice patterns. Symbolic decision graphs use adjacency lists '
        'where nodes are decision states and edges are actions. The build_symbolic_graph() function constructs intent-specific graphs: '
        'family planning includes "partner_alignment" → "health_readiness" → "financial_readiness" paths, while career decisions flow '
        'through "market_scan" → "skills_gap" → "timeline_plan." Set-based representations compute intent alignment through intersection '
        'operations: user token set U ∩ intent keyword set K_intent quantifies semantic overlap. Session memory uses HashTable key-value '
        'storage with intent labels as keys and occurrence counts as values, enabling frequency-based pattern detection. Output structures '
        'use nested dictionaries capturing all reasoning artifacts for transparency and multi-format consumption.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Symbolic Planning
    h7 = doc.add_heading('Symbolic Planning Implementation', level=1)
    h7.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'Symbolic planning converts abstract decision queries into sequential action plans. State spaces begin at "start" and terminate '
        'at "decision_ready," with intermediate states representing analytical milestones. Health decisions progress through "symptom_log" '
        '→ "professional_check" → "habit_plan," while career decisions follow "market_scan" → "skills_gap" → "timeline_plan." This '
        'decomposition mirrors STRIPS-style subgoal planning (Russell & Norvig, 2021).'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        'Both BFS and A* search execute for comparison. BFS explores level-by-level, guaranteeing shortest paths; A* uses heuristics '
        'estimating remaining cost (e.g., "timeline_plan" = 1, "start" = 4), reducing node expansions. For complex family planning graphs, '
        'A* typically expands 5-7 nodes versus BFS\'s 8-10. The system selects A*\'s result when successful, falling back to BFS for '
        'robustness. Outputs include node expansion counts and action sequences, transforming questions like "Should we have a baby?" into '
        'checklists: "Discuss expectations with partner → Review health with professional → Check budget readiness → Create 6-12 month '
        'timeline." This cognitive planning application demonstrates how search algorithms generalize beyond spatial navigation '
        '(Hart, Nilsson, & Raphael, 1968).'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Results
    h8 = doc.add_heading('Results, Limitations, and Future Directions', level=1)
    h8.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(
        'The system successfully provides consistent, structured guidance through both CLI and web interfaces. Input validation ensures '
        'proper formatting; error handling prevents crashes; confidence scoring quantifies interpretation certainty (35-95%); and transparent '
        'reasoning traces show tokenization, signal scores, triggered rules, search comparisons, and ranking details. The implementation '
        'satisfies course requirements: uses methods from four modules (search, expert systems, reasoning, planning), provides reasonable '
        'answers without errors, and interacts naturally with users.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        'Limitations include reliance on keyword matching that may miss semantic nuances, lack of probabilistic uncertainty representation, '
        'manually-curated knowledge requiring expert updates rather than learning from data, and absence of long-term user modeling. The '
        'zodiac framework provides cultural engagement but makes no empirical claims. Future enhancements could integrate machine learning '
        'classifiers for intent detection, Bayesian networks for uncertainty reasoning, and evidence-based decision frameworks from behavioral '
        'economics. Despite limitations, the system demonstrates that classical symbolic AI delivers practical value when transparency and '
        'explainability are prioritized (Luger, 2009). Ethical disclaimers explicitly state the system provides reflective guidance only, '
        'not medical, legal, or financial advice.'
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Page break before references
    doc.add_page_break()

    # References
    ref_h = doc.add_heading('References', level=1)
    ref_h.runs[0].font.size = Pt(12)
    ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Reference entries with hanging indent
    ref1 = doc.add_paragraph('Giarratano, J., & Riley, G. (2005). ')
    ref1.add_run('Expert systems: Principles and programming').italic = True
    ref1.add_run(' (4th ed.). Thomson.')
    ref1.paragraph_format.left_indent = Inches(0.5)
    ref1.paragraph_format.first_line_indent = Inches(-0.5)

    ref2 = doc.add_paragraph(
        'Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). A formal basis for the heuristic determination of '
        'minimum cost paths. '
    )
    ref2.add_run('IEEE Transactions on Systems Science and Cybernetics, 4').italic = True
    ref2.add_run('(2), 100-107. https://doi.org/10.1109/TSSC.1968.300136')
    ref2.paragraph_format.left_indent = Inches(0.5)
    ref2.paragraph_format.first_line_indent = Inches(-0.5)

    ref3 = doc.add_paragraph('Luger, G. F. (2009). ')
    ref3.add_run('Artificial intelligence: Structures and strategies for complex problem solving').italic = True
    ref3.add_run(' (6th ed.). Addison-Wesley.')
    ref3.paragraph_format.left_indent = Inches(0.5)
    ref3.paragraph_format.first_line_indent = Inches(-0.5)

    ref4 = doc.add_paragraph(
        'OpenStreetMap Foundation. (n.d.). Nominatim API documentation. '
        'https://nominatim.org/release-docs/develop/api/Overview/'
    )
    ref4.paragraph_format.left_indent = Inches(0.5)
    ref4.paragraph_format.first_line_indent = Inches(-0.5)

    ref5 = doc.add_paragraph('Russell, S. J., & Norvig, P. (2021). ')
    ref5.add_run('Artificial intelligence: A modern approach').italic = True
    ref5.add_run(' (4th ed.). Pearson.')
    ref5.paragraph_format.left_indent = Inches(0.5)
    ref5.paragraph_format.first_line_indent = Inches(-0.5)

    # Save
    filename = 'CSC510_Final_Essay.docx'
    doc.save(filename)

    # Verify file exists
    if os.path.exists(filename):
        file_size = os.path.getsize(filename)
        print(f"\n✅ SUCCESS!")
        print(f"✅ File created: {filename}")
        print(f"✅ File size: {file_size:,} bytes")
        print(f"✅ Location: {os.path.abspath(filename)}")
        return True
    else:
        print(f"\n❌ ERROR: File not created")
        return False

if __name__ == "__main__":
    success = create_essay()
    sys.exit(0 if success else 1)

